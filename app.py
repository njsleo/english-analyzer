import streamlit as st
import streamlit.components.v1 as components  # 🌟 新增：沙盒组件引擎，用于完美运行发音代码
import json
import pandas as pd
import trafilatura
import io
import datetime
import random
import string
import re
import hashlib
import urllib.parse
import base64
import extra_streamlit_components as esc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls
from openai import OpenAI
from supabase import create_client, Client
from pypdf import PdfReader
from openpyxl.styles import PatternFill, Font, Alignment

# ==========================================
# ⚙️ 核心配置区
# ==========================================
DEEPSEEK_API_KEY = st.secrets["DEEPSEEK_API_KEY"]
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
ADMIN_EMAIL = "75736724@qq.com" # 👑 老板权限
CONTACT_WECHAT = "你的微信号" 

llm_client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ==========================================
# 🍪 商业级加密 Cookie 记忆引擎 
# ==========================================
cookie_manager = esc.CookieManager()

def get_secure_sign(email):
    return hashlib.sha256(f"{email}{SUPABASE_KEY}".encode()).hexdigest()

class SimpleUser:
    def __init__(self, email, uid):
        self.email = email
        self.id = uid

# ==========================================
# 🎨 UI/UX 极致全屏顶导视觉系统 
# ==========================================
st.set_page_config(page_title="顶级英语教研平台-商业版", page_icon="🏛️", layout="wide", initial_sidebar_state="collapsed")

custom_css = """
<style>
    [data-testid="collapsedControl"] { display: none !important; }
    [data-testid="stSidebar"] { display: none !important; }
    
    .block-container { padding-top: 0rem !important; padding-bottom: 1rem !important; margin-top: -1rem !important; max-width: 95% !important;}
    [data-testid="stAppViewBlockContainer"] { padding-top: 0rem !important; }
    [data-testid="stHeader"] { display: none !important; height: 0 !important; }
    .stHeadingContainer { margin-top: -2rem !important; }
    
    h1 { font-size: 1.6rem !important; margin-top: -1rem !important; padding-bottom: 10px !important; }
    h2 { font-size: 1.3rem !important; }
    h3 { font-size: 1.1rem !important; }
    h4 { font-size: 1.05rem !important; }
    h5 { font-size: 1rem !important; }

    .stApp { background-color: #EBF0E5 !important; }
    
    h1, h2, h3, h4, h5 { font-family: 'Times New Roman', 'DengXian', '等线', serif !important; color: #1A1A24; font-weight: bold;}
    p { color: #1A1A24; }
    
    div.row-widget.stRadio > div { flex-direction: row; gap: 8px; flex-wrap: wrap; margin-top: 5px;}
    div[role="radiogroup"] label[data-baseweb="radio"] > div:first-child { display: none !important; } 
    div[role="radiogroup"] label[data-baseweb="radio"] { 
        padding: 4px 16px !important; 
        border-radius: 50px !important; 
        background-color: transparent !important;
        border: 1px solid #C5D1B8 !important;
        cursor: pointer;
        margin: 0 !important;
        transition: all 0.2s ease;
    }
    div[role="radiogroup"] label[data-baseweb="radio"] > div:nth-child(2) { margin-left: 0 !important; width: 100%; text-align: center; }
    div[role="radiogroup"] label[data-baseweb="radio"] p { color: #556070 !important; font-size: 0.85em !important; margin: 0 !important;}
    
    div[role="radiogroup"] label[data-baseweb="radio"]:hover { background-color: #DFE6D8 !important; transform: translateY(-1px); box-shadow: 0 2px 5px rgba(0,0,0,0.05);}
    div[role="radiogroup"] label[data-baseweb="radio"][data-checked="true"] { background-color: #1A1A24 !important; border-color: #1A1A24 !important; box-shadow: 0 4px 10px rgba(0,0,0,0.1); }
    div[role="radiogroup"] label[data-baseweb="radio"][data-checked="true"] p { color: #FFFFFF !important; font-weight: bold !important; }

    div.stButton > button { border-radius: 6px !important; font-weight: 600 !important; border: none !important; box-shadow: 0 2px 4px rgba(0,0,0,0.05); transition: all 0.2s ease; }
    div.stButton > button:hover { transform: translateY(-1px); box-shadow: 0 4px 8px rgba(0,0,0,0.1); }
    .stTextInput input, .stTextArea textarea, .stSelectbox > div > div { border-radius: 6px !important; border: 1px solid #D8DFD0 !important; background-color: #F5F7EC !important; color: #2C3E50 !important;}
    
    div[data-baseweb="tab-list"] { gap: 6px; padding-bottom: 5px; }
    div[data-baseweb="tab"] { padding: 8px 16px !important; font-size: 0.9em !important; border-radius: 6px 6px 0 0; background-color: transparent; }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# ==========================================
# 🛠️ 核心工具函数区
# ==========================================
def set_font(run, ascii_font='Times New Roman', east_asia_font='等线'): 
    run.font.name = ascii_font; run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)

def export_plain_text_to_word(text_content):
    doc = Document(); style = doc.styles['Normal']; style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '等线'); doc.add_heading('📖 英语精读教案 (备份)', 0)
    for line in text_content.split('\n'):
        if line.strip(): p = doc.add_paragraph(); set_font(p.add_run(line))
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def generate_beautiful_word(analysis_data, full_text=""):
    doc = Document()
    try:
        bg = parse_xml(f'<w:background {nsdecls("w")} w:color="F4F6F1"/>')
        doc.settings.element.insert(0, bg); shape = parse_xml(f'<w:displayBackgroundShape {nsdecls("w")}/>'); doc.settings.element.append(shape)
    except: pass
    style = doc.styles['Normal']; style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    doc.add_heading('📖 英语专家级精读教案', level=1)
    if full_text:
        doc.add_heading('一、 英文原文', level=2)
        run_org = doc.add_paragraph().add_run(full_text.strip()); run_org.font.size, run_org.font.color.rgb = Pt(11), RGBColor(0x33, 0x33, 0x33); set_font(run_org)
    doc.add_heading('二、 逐句解析', level=2)
    for i, s in enumerate(analysis_data.get('sentences', [])):
        run_en = doc.add_paragraph().add_run(f"[{i+1}] {s.get('en', '')}"); run_en.bold, run_en.font.size = True, Pt(12); set_font(run_en)
        run_cn = doc.add_paragraph().add_run(f"译文：{s.get('cn', '')}"); run_cn.font.size, run_cn.font.color.rgb = Pt(10.5), RGBColor(0x55, 0x55, 0x55); set_font(run_cn)
        p_syn = doc.add_paragraph(); p_syn.paragraph_format.left_indent = Pt(15)
        r_syn = p_syn.add_run(f"🔍 语法：{s.get('syntax','')}\n💡 词法：{s.get('words','')}"); r_syn.font.size = Pt(10.5); set_font(r_syn)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    v_list = analysis_data.get('core_vocabulary', [])
    if v_list:
        doc.add_heading('三、 核心词汇表', level=2)
        table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
        for i, h in enumerate(['单词', '音标', '释义', '例句']): set_font(table.rows[0].cells[i].paragraphs[0].add_run(h))
        for v in v_list:
            row = table.add_row().cells; row[0].text, row[1].text, row[2].text, row[3].text = v.get('word',''), v.get('phonetic',''), v.get('translation',''), v.get('usage_examples','')
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

def fetch_text_smart(url): 
    try: return trafilatura.extract(trafilatura.fetch_url(url)) if trafilatura.fetch_url(url) else "⚠️ 未能识别正文"
    except: return "抓取异常"

def extract_text_from_file(uploaded_file):
    if uploaded_file.type == "text/plain": return uploaded_file.read().decode("utf-8")
    elif uploaded_file.type == "application/pdf":
        return "\n".join([page.extract_text() for page in PdfReader(uploaded_file).pages if page.extract_text()])
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return "\n".join([para.text for para in Document(uploaded_file).paragraphs])
    return ""

def format_reading_text(text):
    cleaned = re.sub(r'\n\s*\n', '§§§', text)
    cleaned = cleaned.replace('\n', ' ')
    paragraphs = [p.strip() for p in cleaned.split('§§§') if p.strip()]
    html = ""
    for p in paragraphs: html += f"<div style='margin-bottom: 10px;'>{p}</div>"
    return html

def export_styled_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='我的生词本')
        worksheet = writer.sheets['我的生词本']
        header_fill = PatternFill(start_color="D3DCCB", end_color="D3DCCB", fill_type="solid")
        row_fill = PatternFill(start_color="F5F7EC", end_color="F5F7EC", fill_type="solid")
        header_font = Font(name="等线", bold=True, color="1F4E79", size=12)
        base_font = Font(name="等线", size=11, color="2C3E50")
        align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
        align_left = Alignment(horizontal="left", vertical="center", wrap_text=True)

        worksheet.row_dimensions[1].height = 30
        for col_num, value in enumerate(df.columns.values):
            cell = worksheet.cell(row=1, column=col_num + 1)
            cell.fill = header_fill; cell.font = header_font; cell.alignment = align_center
        for row_num in range(2, len(df) + 2):
            worksheet.row_dimensions[row_num].height = 40
            for col_num in range(1, len(df.columns) + 1):
                cell = worksheet.cell(row=row_num, column=col_num)
                cell.fill = row_fill; cell.font = base_font
                if col_num in [1, 2, 4]: cell.alignment = align_center
                else: cell.alignment = align_left
        col_widths = {'A': 16, 'B': 18, 'C': 35, 'D': 10, 'E': 45, 'F': 60}
        for col, width in col_widths.items(): worksheet.column_dimensions[col].width = width
    return output.getvalue()

# 🌟 查词沙盒渲染器 (解决防注入不发音问题)
def render_dictionary_card(word_data):
    safe_word = urllib.parse.quote(word_data.get('word', '')).replace("'", "%27")
    audio_url = f"https://dict.youdao.com/dictvoice?audio={safe_word}&type=2"
    dict_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body {{ margin: 0; font-family: "Times New Roman", "等线", serif; background-color: transparent; }}
        .card {{ background-color:#F5F7EC; padding:15px; border-radius:6px; border:1px solid #D8DFD0; color: #2C3E50; }}
        .audio-btn {{ cursor: pointer; font-size: 1.15em; margin-left: 5px; transition: transform 0.2s ease; display: inline-block; }}
        .audio-btn:hover {{ transform: scale(1.3); }}
    </style>
    </head>
    <body>
        <div class="card">
            <strong style='font-size: 1.15em; color: #1A1A24;'>{word_data.get('word')}</strong> 
            <span style='color: #666; margin-left: 5px;'>{word_data.get('phonetic')}</span> 
            <span class="audio-btn" onclick="new Audio('{audio_url}').play()" title="点击听纯正发音">🔊</span><br><br>
            <strong>释义：</strong>{word_data.get('translation')}<br><br>
            <strong>记忆：</strong><span style='color: #555;'>{word_data.get('memory_tip')}</span>
        </div>
    </body>
    </html>
    """
    components.html(dict_html, height=160)

# 🌟 词库表格沙盒渲染器 (解决防注入不发音问题)
def render_vocabulary_table(df):
    html_table = """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body { margin: 0; font-family: "Times New Roman", "等线", serif; background-color: transparent; }
        .table-wrapper { background-color: #F5F7EC; border: 1px solid #D8DFD0; border-radius: 8px; overflow: auto; height: 580px; box-shadow: 0 4px 15px rgba(0,0,0,0.03); }
        table { width: 100%; border-collapse: collapse; text-align: left; }
        th { padding: 12px 16px; border-bottom: 1px solid #D8DFD0; color: #1F4E79; position: sticky; top: 0; background-color: #DFE6D8; z-index: 1; font-weight: bold; }
        td { padding: 12px 16px; border-bottom: 1px solid #EAECEF; color: #2C3E50; }
        .audio-btn { cursor: pointer; margin-left: 8px; font-size: 1.15em; transition: transform 0.2s ease; display: inline-block; }
        .audio-btn:hover { transform: scale(1.3); text-shadow: 0 2px 4px rgba(0,0,0,0.15); }
        .tag { background-color:#D3DCCB; padding:3px 8px; border-radius:4px; font-size:0.85em; color:#111; }
    </style>
    </head>
    <body>
    <div class="table-wrapper">
    <table>
        <thead>
            <tr><th>单词</th><th>音标</th><th>释义</th><th>级别</th><th>记忆法</th><th>实用例句</th></tr>
        </thead>
        <tbody>
    """
    for _, row in df.iterrows():
        safe_word = urllib.parse.quote(row.get('word', '')).replace("'", "%27")
        audio_link = f"https://dict.youdao.com/dictvoice?audio={safe_word}&type=2"
        html_table += f"<tr><td style='font-weight: bold; color: #1A1A24; font-size: 1.1em;'>{row.get('word','')}</td><td style='color: #666; white-space: nowrap;'>{row.get('phonetic','')}<span class='audio-btn' onclick=\"new Audio('{audio_link}').play()\" title='点击听发音'>🔊</span></td><td>{row.get('translation','')}</td><td><span class='tag'>{row.get('tags','')}</span></td><td style='color: #555;'>{row.get('memory_tip','')}</td><td style='color: #444; font-size: 0.9em;'>{row.get('usage_examples','')}</td></tr>"
    html_table += "</tbody></table></div></body></html>"
    components.html(html_table, height=600)

# ==========================================
# 🔐 认证与无感登录系统
# ==========================================
if 'user' not in st.session_state: st.session_state['user'] = None

if st.session_state['user'] is None:
    c_email = cookie_manager.get("saved_email")
    c_uid = cookie_manager.get("saved_uid")
    c_sign = cookie_manager.get("saved_sign")
    if c_email and c_uid and c_sign:
        if c_sign == get_secure_sign(c_email): st.session_state['user'] = SimpleUser(c_email, c_uid)

if st.session_state['user'] is None:
    st.markdown("<h1 style='text-align: center; margin-top:50px;'>🏛️ 顶级英语精读工作台</h1>", unsafe_allow_html=True)
    _, col_auth, _ = st.columns([1, 1, 1])
    with col_auth:
        tab_login, tab_signup = st.tabs(["🔑 登录", "🎟️ 凭邀请码注册"])
        with tab_login:
            email = st.text_input("邮箱"); pwd = st.text_input("密码", type="password")
            if st.button("进入系统", use_container_width=True, type="primary"):
                try: 
                    res = supabase.auth.sign_in_with_password({"email": email, "password": pwd})
                    st.session_state['user'] = res.user
                    cookie_manager.set("saved_email", res.user.email, max_age=30*24*3600)
                    cookie_manager.set("saved_uid", res.user.id, max_age=30*24*3600)
                    cookie_manager.set("saved_sign", get_secure_sign(res.user.email), max_age=30*24*3600)
                    st.rerun()
                except: st.error("账号或密码有误")
        with tab_signup:
            s_email = st.text_input("设置邮箱"); s_pwd = st.text_input("设置密码(>6位)", type="password"); s_code = st.text_input("邀请码")
            if st.button("注册"):
                code_res = supabase.table('invitation_codes').select('*').eq('code', s_code).eq('is_used', False).execute()
                if code_res.data:
                    try:
                        supabase.auth.sign_up({"email": s_email, "password": s_pwd})
                        exp = (datetime.datetime.now() + datetime.timedelta(days=code_res.data[0]['duration_days'])).isoformat()
                        supabase.table('invitation_codes').update({'is_used': True, 'used_by': s_email}).eq('code', s_code).execute()
                        supabase.table('subscriptions').insert({'user_email': s_email, 'expires_at': exp, 'role': 'user'}).execute()
                        st.success("注册成功！请切换登录。")
                    except: st.error("注册失败，可能邮箱已被使用。")
                else: st.error("邀请码无效或已使用")
    st.stop()

# ==========================================
# 🛡️ 订阅与 RBAC 权限系统
# ==========================================
USER_EMAIL = st.session_state['user'].email; CURRENT_USER_ID = st.session_state['user'].id
IS_SUPER_ADMIN = (USER_EMAIL == ADMIN_EMAIL) 

current_exp = None
is_expired = False
user_role = "user"

sub_res = supabase.table('subscriptions').select('*').eq('user_email', USER_EMAIL).execute()
if sub_res.data:
    current_exp = datetime.datetime.fromisoformat(sub_res.data[0]['expires_at'])
    if datetime.datetime.now() > current_exp and not IS_SUPER_ADMIN: is_expired = True
    user_role = sub_res.data[0].get('role', 'user')
else:
    if not IS_SUPER_ADMIN: is_expired = True

IS_ADMIN = IS_SUPER_ADMIN or (user_role == 'admin')

if not IS_SUPER_ADMIN and is_expired:
    st.warning("⚠️ 您的 VIP 授权已到期，系统已暂停您的操作权限。")
    st.info(f"👉 您的账号资料已安全锁定。请联系管理员微信 **{CONTACT_WECHAT}** 进行续费激活，解锁全部权限！")
    if st.button("🚪 退出系统"): 
        cookie_manager.delete("saved_email"); cookie_manager.delete("saved_uid"); cookie_manager.delete("saved_sign")
        st.session_state['user'] = None; st.rerun()
    st.stop()

# ==========================================
# 🌟 全局导航栏 
# ==========================================
menu_options = ["📚 公共教材图书馆", "🔍 智能精读教研室", "🗂️ 文章分类档案馆", "🔠 词库与大纲"]
if IS_SUPER_ADMIN: menu_options.append("👑 创始人控制台") 

if 'nav_page' not in st.session_state: st.session_state['nav_page'] = "📚 公共教材图书馆"
default_idx = menu_options.index(st.session_state['nav_page']) if st.session_state['nav_page'] in menu_options else 0

col_nav, col_info, col_logout = st.columns([6, 2.5, 0.8], gap="medium")

with col_nav:
    page = st.radio("主导航", menu_options, index=default_idx, horizontal=True, label_visibility="collapsed")
    st.session_state['nav_page'] = page 

with col_info:
    role_badge = "👑 馆长" if IS_ADMIN else "👤 会员"
    status_icon = "🔴" if is_expired else "🟢"
    exp_text = current_exp.strftime('%Y-%m-%d') if current_exp else "终身"
    st.markdown(f"<div style='text-align: right; padding-top: 15px; color: #556070; font-size: 0.85em;'>{role_badge} <b>{USER_EMAIL}</b> &nbsp;|&nbsp; {status_icon} {exp_text}</div>", unsafe_allow_html=True)

with col_logout:
    st.write("")
    if st.button("🚪 退出", use_container_width=True): 
        cookie_manager.delete("saved_email"); cookie_manager.delete("saved_uid"); cookie_manager.delete("saved_sign")
        st.session_state['user'] = None; st.rerun()

st.markdown("<hr style='margin-top: 0px; margin-bottom: 15px; border: 0; border-top: 1px solid #D8DFD0;'>", unsafe_allow_html=True)

# ==========================================
# 👑 模块：创始人控制台
# ==========================================
if IS_SUPER_ADMIN and page == "👑 创始人控制台":
    tab_gen, tab_users, tab_codes = st.tabs(["🎟️ 激活码生成", "👥 用户管理 & 授权", "📋 激活码查账明细"])
    
    with tab_gen:
        st.markdown("#### 🔨 生成新激活码")
        with st.form("gen_code_form"):
            plan = st.radio("授权时长：", ["7天试用", "1个月", "3个月", "1年", "终身"], horizontal=True)
            days_map = {"7天试用": 7, "1个月": 30, "3个月": 90, "1年": 365, "终身": 36500}
            if st.form_submit_button("🔨 立即生成", type="primary"):
                new_code = f"VIP-{''.join(random.choices(string.ascii_uppercase + string.digits, k=8))}"
                try:
                    supabase.table('invitation_codes').insert({"code": new_code, "duration_days": days_map[plan], "is_used": False}).execute()
                    st.success(f"生成成功: {new_code}"); st.code(new_code)
                except: st.error("生成失败")

    with tab_users:
        st.markdown("#### 👥 客户管理 & 权限分配")
        try:
            sub_data = supabase.table('subscriptions').select('*').execute().data
            if sub_data:
                df_subs = pd.DataFrame(sub_data); now_dt = datetime.datetime.now(); df_subs['到期时间'] = pd.to_datetime(df_subs['expires_at'])
                df_subs['状态'] = df_subs['到期时间'].apply(lambda x: "🔴 已过期" if x < now_dt else "🟢 正常")
                st.metric("总注册用户数", len(df_subs))
                selected_user = st.selectbox("🔍 搜索或选择要操作的客户账号：", df_subs['user_email'].tolist())
                
                if selected_user:
                    user_info = df_subs[df_subs['user_email'] == selected_user].iloc[0]; curr_exp = user_info['到期时间']
                    curr_role = user_info.get('role', 'user')
                    st.markdown(f"<div style='background:#F5F7EC; padding:15px; border-radius:8px; border:1px solid #D8DFD0; margin-bottom:15px;'><b style='font-size:1.1em;'>客户：{selected_user}</b><br>当前状态：{user_info['状态']}<br>系统角色：{curr_role}<br>到期时间：{curr_exp.strftime('%Y-%m-%d %H:%M:%S')}</div>", unsafe_allow_html=True)
                    
                    c1, c2 = st.columns(2)
                    with c1:
                        st.markdown("##### ⚡ 一键充值 (免密续费)")
                        col_r1, col_r2 = st.columns(2); add_days = 0
                        if col_r1.button("💸 续费 30 天", use_container_width=True): add_days = 30
                        if col_r2.button("💸 续费 365 天", use_container_width=True): add_days = 365
                        if add_days > 0:
                            base_date = curr_exp if curr_exp > now_dt else now_dt
                            new_exp = base_date + datetime.timedelta(days=add_days)
                            try:
                                supabase.table('subscriptions').update({'expires_at': new_exp.isoformat()}).eq('user_email', selected_user).execute()
                                st.success(f"✅ 续费成功！已为 {selected_user} 增加 {add_days} 天。"); st.rerun()
                            except: st.error("续费失败")
                    with c2:
                        st.markdown("##### 🛡️ 系统角色授权")
                        if curr_role == 'user':
                            if st.button("👑 提拔为【内容管理员】", type="primary", use_container_width=True):
                                supabase.table('subscriptions').update({'role': 'admin'}).eq('user_email', selected_user).execute()
                                st.success(f"✅ 已将 {selected_user} 设为管理员！"); st.rerun()
                        else:
                            if st.button("⬇️ 降级为【普通用户】", use_container_width=True):
                                supabase.table('subscriptions').update({'role': 'user'}).eq('user_email', selected_user).execute()
                                st.success(f"✅ 已取消 {selected_user} 的管理员权限！"); st.rerun()
            else: st.info("当前还没有注册用户。")
        except: pass
        
    with tab_codes:
        st.markdown("#### 📋 激活码核销账本")
        try:
            codes_data = supabase.table('invitation_codes').select('*').execute().data
            if codes_data:
                df_codes = pd.DataFrame(codes_data); df_codes['状态'] = df_codes['is_used'].apply(lambda x: "🔴 已核销" if x else "🟢 未使用")
                if 'used_by' in df_codes.columns:
                    display_codes = df_codes[['code', 'duration_days', '状态', 'used_by', 'created_at']]
                    display_codes.columns = ['激活码', '授权天数', '状态', '使用者', '生成时间']; display_codes['使用者'] = display_codes['使用者'].fillna('-')
                else:
                    display_codes = df_codes[['code', 'duration_days', '状态', 'created_at']]
                    display_codes.columns = ['激活码', '授权天数', '状态', '生成时间']
                st.dataframe(display_codes.sort_values(by='生成时间', ascending=False), use_container_width=True, hide_index=True)
        except: pass

# ==========================================
# 📚 模块：公共教材图书馆
# ==========================================
elif page == "📚 公共教材图书馆":
    
    if 'reading_book_title' not in st.session_state:
        st.session_state['reading_book_title'] = None

    base_categories = ["全部", "新概念", "小学教材", "初中教材", "高中教材", "大学四六级", "雅思托福", "英文名著", "外刊新闻", "冒险悬疑", "科幻奇幻", "浪漫爱情", "历史传记", "童话寓言", "短篇小说", "商业科技", "喜剧戏剧", "影视原著", "课外阅读", "其他"]
    
    lib_data = []
    try:
        lib_data_raw = supabase.table('public_library').select('*').execute().data
        lib_data = [a for a in lib_data_raw if a.get('category') != "公共词库"] if lib_data_raw else []
    except: pass

    df_lib = pd.DataFrame(lib_data)
    if not df_lib.empty:
        db_cats = list(df_lib['category'].dropna().unique())
        final_categories = [c for c in base_categories if c == "全部" or c in db_cats] + [c for c in db_cats if c not in base_categories]
    else: final_categories = ["全部"]

    if st.session_state['reading_book_title'] is None:
        if IS_ADMIN:
            with st.expander("👑 馆长专属：上传新教材/小说", expanded=False):
                lib_title = st.text_input("篇目标题"); lib_cat = st.selectbox("选择分类", base_categories[1:])
                cover_file = st.file_uploader("🖼️ 上传自定义封面 (可选, 不传则自动生成绝美图片)", type=["png", "jpg", "jpeg"])
                upload_method = st.radio("正文录入方式", ["手动粘贴", "📂 上传本地文档"], horizontal=True, label_visibility="collapsed")
                lib_content = st.text_area("正文", height=100) if upload_method == "手动粘贴" else ""
                if upload_method != "手动粘贴":
                    uploaded_file = st.file_uploader("选择纯文本/PDF/Word文档", type=["pdf", "docx", "txt"])
                    if uploaded_file: lib_content = extract_text_from_file(uploaded_file); st.success("提取成功！")
                
                if st.button("⬆️ 上传至公共书架", type="primary"):
                    if lib_title and lib_content.strip():
                        cover_b64 = ""
                        if cover_file: cover_b64 = "data:image/jpeg;base64," + base64.b64encode(cover_file.read()).decode()
                        try: supabase.table('public_library').insert({"title": lib_title, "category": lib_cat, "content": lib_content, "cover_image": cover_b64}).execute()
                        except: supabase.table('public_library').insert({"title": lib_title, "category": lib_cat, "content": lib_content}).execute()
                        st.success("✅ 上传成功！"); st.session_state['reading_book_title'] = None; st.rerun()

        cat_filter = st.radio("分类", final_categories, horizontal=True, label_visibility="collapsed", key="cat_radio")
        st.write("---")

        filtered_lib = [a for a in lib_data if a.get('category') == cat_filter] if cat_filter != "全部" else lib_data
        
        if filtered_lib:
            cols = st.columns(6)
            for i, book in enumerate(filtered_lib):
                with cols[i % 6]:
                    cover_img_src = book.get('cover_image')
                    if not cover_img_src:
                        title_hash = hashlib.md5(book['title'].encode()).hexdigest()[:8]
                        cover_img_src = f'https://picsum.photos/seed/{title_hash}/400/550'
                    
                    tag_color = "#FF4B4B" if i % 3 == 0 else ("#00B4D8" if i % 2 == 0 else "#FFB703") 
                    
                    card_html = f"""
                    <div style='background-color: #fff; border-radius: 8px; padding: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.05); margin-bottom: -15px;'>
                        <div style='position: relative; width: 100%; padding-top: 140%; border-radius: 4px; overflow: hidden; background: #eee;'>
                            <img src='{cover_img_src}' style='position: absolute; top: 0; left: 0; width: 100%; height: 100%; object-fit: cover;'>
                            <div style='position: absolute; bottom: 15px; left: 0; background-color: {tag_color}; color: white; font-size: 0.75rem; font-weight: bold; padding: 8px 4px; writing-mode: vertical-lr; text-orientation: upright; letter-spacing: 2px; border-radius: 0 4px 4px 0; box-shadow: 2px 0 5px rgba(0,0,0,0.2);'>
                                {book.get('category')}
                            </div>
                        </div>
                        <div style='margin-top: 10px; margin-bottom: 12px; text-align: center;'>
                            <div style='font-weight: bold; font-size: 0.9em; color: #1A1A24; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; font-family: "Times New Roman", serif;'>{book.get('title')}</div>
                            <div style='font-size: 0.7em; color: #8892B0; margin-top: 2px;'>Word Count: {len(book.get('content','').split())}</div>
                        </div>
                    </div>
                    """
                    st.markdown(card_html, unsafe_allow_html=True)
                    if st.button("📖 立即阅读", key=f"read_{book['id']}", use_container_width=True):
                        st.session_state['reading_book_title'] = book['title']
                        st.rerun()
        else:
            st.info("💡 当前分类下暂无教材，等待馆长上新！")

    else:
        selected_lib_item = next((b for b in lib_data if b['title'] == st.session_state['reading_book_title']), None)
        
        if selected_lib_item:
            c_back, c_space, c_font = st.columns([1.5, 5, 2.5])
            with c_back:
                if st.button("⬅️ 返回书架", type="primary"):
                    st.session_state['reading_book_title'] = None
                    st.rerun()
            with c_font:
                font_opt = st.radio("阅读字号", ["标准", "放大", "特大"], horizontal=True, index=0, label_visibility="collapsed")
                font_map = {"标准": "1.05em", "放大": "1.25em", "特大": "1.45em"}
                current_fs = font_map[font_opt]
            
            col_read, col_tools = st.columns([4.5, 1.5], gap="large")
            
            with col_read:
                st.markdown(f"#### {selected_lib_item.get('title')}")
                clean_html_text = format_reading_text(selected_lib_item.get('content', ''))
                paper_bg = "#F5F7EC" 
                st.markdown(f"<div style='background-color: {paper_bg}; padding: 40px 60px; border-radius: 8px; font-family: \"Times New Roman\", serif; font-size: {current_fs}; color: #2C3E50; line-height: 1.8; text-align: justify; height: 75vh; overflow-y: auto; border: 1px solid #D8DFD0; box-shadow: 0 4px 15px rgba(0,0,0,0.03); transition: font-size 0.3s ease;'>{clean_html_text}</div>", unsafe_allow_html=True)
            
            with col_tools:
                st.markdown("#### 🛠️ 伴读助手")
                tab_dict, tab_clip = st.tabs(["🔍 查词", "📝 摘抄"])
                with tab_dict:
                    st.caption("复制左侧生词粘贴查阅")
                    lookup_word = st.text_input("输入英文生词", label_visibility="collapsed", placeholder="例如: consecutive", key="lib_word_input")
                    if st.button("💡 翻译并存库", type="primary", use_container_width=True, key="lib_btn_trans"):
                        if lookup_word:
                            with st.spinner("查词中..."):
                                prompt = f"""分析单词: {lookup_word}。返回纯JSON: {{"word":"{lookup_word}","phonetic":"音标","translation":"精准中文释义","memory_tip":"一句精简的词根或联想记忆法","usage_examples":"一个简短实用的英文例句及中文","tags":"阅读生词"}}"""
                                try:
                                    res = llm_client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], response_format={"type":"json_object"})
                                    word_data = json.loads(res.choices[0].message.content)
                                    # 🌟 修复发音沙盒调用
                                    render_dictionary_card(word_data)
                                    word_data['user_id'] = CURRENT_USER_ID; supabase.table('vocabulary').insert(word_data).execute(); st.success("✅ 已存入记忆库")
                                except: st.error("查词失败")
                with tab_clip:
                    st.caption("复制左侧难句解析")
                    clip_sentence = st.text_area("输入句子", label_visibility="collapsed", height=100, placeholder="粘贴想精读的句子...", key="lib_clip_input")
                    if st.button("✍️ 解析并归档", type="primary", use_container_width=True, key="lib_btn_clip"):
                        if clip_sentence:
                            with st.spinner("解析中..."):
                                prompt = f"""深度解析此句，返回JSON: {{"sentences":[{{"en":"{clip_sentence}","cn":"精美的翻译","syntax":"极简语法框架拆解","words":"核心词组解析"}}]}}"""
                                try:
                                    res = llm_client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], response_format={"type":"json_object"})
                                    clip_data = json.loads(res.choices[0].message.content); s = clip_data['sentences'][0]
                                    txt = f"[{1}] {s.get('en','')}\n译：{s.get('cn','')}\n🔍 语法：{s.get('syntax','')}\n💡 词法：{s.get('words','')}\n\n"
                                    supabase.table('articles').insert({"user_id": CURRENT_USER_ID, "content": clip_sentence, "teaching_plan": txt, "translation": json.dumps(clip_data), "category": "摘抄好句"}).execute()
                                    st.success("✅ 已存至档案馆")
                                    st.markdown(f"<div style='font-size:0.9em; background:#F5F7EC; padding:10px; border-radius:5px; border:1px solid #D8DFD0;'><b>译：</b>{s.get('cn')}<br><br><b>语法：</b>{s.get('syntax')}</div>", unsafe_allow_html=True)
                                except: st.error("解析失败")

# ==========================================
# 🔍 模块：教研室
# ==========================================
elif page == "🔍 智能精读教研室":
    col1, col2 = st.columns([4, 1])
    with col1: url = st.text_input("🔗 输入英文文章链接：")
    with col2: 
        st.write(""); st.write("")
        if st.button("🛰️ 提取网页", use_container_width=True):
            if url: st.session_state['temp_text'] = fetch_text_smart(url)
    
    final_text = st.text_area("📝 待分析文本：", value=st.session_state.get('temp_text', ""), height=200)
    
    if st.button("🧠 生成专家级教案", type="primary"):
        if not final_text.strip(): st.error("请先输入文本")
        else:
            with st.spinner("AI正在切片..."):
                prompt = f"""以JSON格式输出全句拆解：{{"sentences": [{{"en": "原句英文", "cn": "翻译", "syntax": "极简语法", "words": "核心词法"}}], "core_vocabulary": [{{"word": "单词", "phonetic": "音标", "translation": "释义", "memory_tip": "记忆法", "usage_examples": "造句", "tags": "级别"}}]}} 待分析：\n{final_text[:5000]}""" 
                try:
                    res = llm_client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], response_format={"type":"json_object"})
                    st.session_state['analysis_result'] = json.loads(res.choices[0].message.content); st.session_state['article_content'] = final_text; st.rerun()
                except Exception: st.error("分析失败")

    if 'analysis_result' in st.session_state:
        res = st.session_state['analysis_result']; st.divider()
        c1, c2, c3 = st.columns(3)
        with c1: st.download_button("📝 导出 Word", data=generate_beautiful_word(res, st.session_state.get('article_content', '')), file_name="教案.docx", use_container_width=True)
        with c2: cat = st.selectbox("📂 保存分类：", ["精读课文", "课外拓展", "考试阅读", "未分类"], label_visibility="collapsed")
        with c3:
            if st.button("☁️ 归档至私人空间", use_container_width=True):
                txt = "".join([f"[{i+1}] {s.get('en','')}\n译：{s.get('cn','')}\n🔍 语法：{s.get('syntax','').replace('*', '')}\n💡 词法：{s.get('words','').replace('*', '')}\n\n" for i,s in enumerate(res.get('sentences', []))])
                try:
                    supabase.table('articles').insert({"user_id": CURRENT_USER_ID, "content": st.session_state['article_content'], "teaching_plan": txt, "translation": json.dumps(res), "category": cat}).execute()
                    for v in res.get('core_vocabulary', []): v["user_id"] = CURRENT_USER_ID; supabase.table('vocabulary').insert(v).execute()
                    st.success("✅ 归档成功！")
                except Exception: st.error("保存失败")
                    
        for i, s in enumerate(res.get('sentences', [])):
            st.markdown(f"""<div style='background:#F5F7EC; border-radius:8px; padding:12px; margin-bottom:8px; border:1px solid #D8DFD0;'>
                <div style='font-family: Times New Roman; font-size:1.05em; font-weight:bold;'>[{i+1}] {s.get('en','')}</div><div style='color:#555; font-size:0.95em;'>译：{s.get('cn','')}</div>
                <div style='font-size:0.9em; margin-top:4px;'><span style='color:#1F4E79;'>🔍 语法：</span>{s.get('syntax','')}</div><div style='font-size:0.9em;'><span style='color:#C00000;'>💡 词法：</span>{s.get('words','')}</div></div>""", unsafe_allow_html=True)

# ==========================================
# 🗂️ 档案馆
# ==========================================
elif page == "🗂️ 文章分类档案馆":
    try:
        arts_data = supabase.table('articles').select('*').eq('user_id', CURRENT_USER_ID).execute().data
        if arts_data:
            df_arts = pd.DataFrame(arts_data); categories = ["全部"] + list(df_arts['category'].dropna().unique())
            tabs = st.tabs(categories)
            for i, tab in enumerate(tabs):
                with tab:
                    filtered_arts = [a for a in arts_data if a.get('category') == categories[i]] if categories[i] != "全部" else arts_data
                    if filtered_arts:
                        col_list, col_content = st.columns([1, 4], gap="large")
                        with col_list:
                            options = [f"{idx+1}. {a.get('content', '')[:25]}..." for idx, a in enumerate(filtered_arts)]
                            selected_title = st.radio("选择文章", options, key=f"radio_{i}", label_visibility="collapsed")
                        with col_content:
                            selected_art = filtered_arts[options.index(selected_title)]
                            art_id = selected_art.get('id'); raw_json = selected_art.get('translation', '')
                            try: full_analysis = json.loads(raw_json) if raw_json else None
                            except: full_analysis = None
                            
                            c1, c2 = st.columns(2)
                            with c1: 
                                word_data = generate_beautiful_word(full_analysis, selected_art.get('content', '')) if full_analysis else export_plain_text_to_word(selected_art.get('teaching_plan', ''))
                                st.download_button("📥 重新导出", data=word_data, file_name="归档教案.docx", use_container_width=True, key=f"dl_{art_id}_{i}")
                            with c2: 
                                if st.button("🗑️ 永久删除", key=f"del_{art_id}_{i}", use_container_width=True):
                                    supabase.table('articles').delete().eq('id', art_id).execute(); st.rerun()
                                    
                            st.markdown("##### 📰 原文/摘抄"); st.markdown(f"<div style='background-color:#F5F7EC; padding:12px; border-radius:6px; border:1px solid #D8DFD0; max-height:120px; overflow-y:auto; margin-bottom:15px;'>{selected_art.get('content','')}</div>", unsafe_allow_html=True)
                            st.markdown("##### 🔬 解析"); st.markdown(f"<div style='background-color:#F5F7EC; padding:16px; border-radius:6px; border:1px solid #D8DFD0; white-space:pre-wrap;'>{selected_art.get('teaching_plan','').strip()}</div>", unsafe_allow_html=True)
                    else: st.info("暂无记录。")
        else: st.info("空空如也。")
    except: pass

# ==========================================
# 🔠 模块：词库与大纲 (🌟 沙盒发音表格完美修复)
# ==========================================
elif page == "🔠 词库与大纲":
    tab_mine, tab_public = st.tabs(["📓 我的私人生词本", "🌍 公共大纲词库"])
    
    with tab_mine:
        try:
            vocab_data = supabase.table('vocabulary').select('*').eq('user_id', CURRENT_USER_ID).execute().data
            if vocab_data:
                df_vocab = pd.DataFrame(vocab_data)
                
                col_m1, col_m2 = st.columns([1, 1])
                with col_m1: st.metric("生词量", len(df_vocab))
                with col_m2: 
                    st.write("")
                    manage_mode = st.toggle("🛠️ 开启打钩/批量管理模式")
                
                if manage_mode:
                    st.info("💡 请在表格第一列打钩（勾选）你要操作的单词。")
                    df_manage = df_vocab[['word', 'phonetic', 'translation', 'tags', 'memory_tip', 'usage_examples']].copy()
                    df_manage.insert(0, "☑️ 勾选", False)
                    
                    edited_df = st.data_editor(
                        df_manage,
                        hide_index=True,
                        use_container_width=True,
                        column_config={"☑️ 勾选": st.column_config.CheckboxColumn("☑️ 勾选", default=False, width="small")}
                    )
                    
                    selected_df = edited_df[edited_df["☑️ 勾选"] == True]
                    
                    st.write("---")
                    c1, c2, c3 = st.columns(3)
                    
                    export_cols_map = {
                        'word': '单词', 
                        'phonetic': '音标', 
                        'translation': '释义', 
                        'tags': '级别', 
                        'memory_tip': '记忆法', 
                        'usage_examples': '例句'
                    }
                    
                    df_export_all = df_vocab[['word', 'phonetic', 'translation', 'tags', 'memory_tip', 'usage_examples']].rename(columns=export_cols_map)
                    excel_all = export_styled_excel(df_export_all)
                    c1.download_button("📥 导出【全部】生词本", excel_all, "我的全部生词本.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                    
                    if not selected_df.empty:
                        df_export_sel = selected_df.drop(columns=['☑️ 勾选']).rename(columns=export_cols_map)
                        excel_sel = export_styled_excel(df_export_sel)
                        c2.download_button(f"📥 导出选中的 {len(selected_df)} 个", excel_sel, "选中的生词.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
                        
                        if c3.button(f"🗑️ 删除选中的 {len(selected_df)} 个", type="primary", use_container_width=True):
                            for w in selected_df['word']:
                                supabase.table('vocabulary').delete().eq('user_id', CURRENT_USER_ID).eq('word', w).execute()
                            st.success("✅ 删除成功！"); st.rerun()
                    else:
                        c2.button("📥 导出选中的 (请先打钩)", disabled=True, use_container_width=True)
                        c3.button("🗑️ 删除选中的 (请先打钩)", disabled=True, use_container_width=True)

                else:
                    cat_filter = st.radio("🎓 分类筛选", ["全部"] + list(df_vocab['tags'].dropna().unique()), horizontal=True, label_visibility="collapsed")
                    display_df = df_vocab[df_vocab['tags'] == cat_filter] if cat_filter != "全部" else df_vocab
                    
                    # 🌟 核心：调用独立渲染函数，保证 HTML 完全隔离不受干扰，确保点击能够发音
                    render_vocabulary_table(display_df)
                
            else: st.info("📓 词汇库还是空的，快去阅读文章添加生词吧！")
        except Exception as e: pass

    with tab_public:
        if IS_ADMIN:
            with st.expander("👑 馆长专属：用 AI 批量生成公共词库", expanded=False):
                st.info("💡 支持直接粘贴单词，或上传纯文本(txt/docx)。AI会自动解析并上架。")
                base_categories = ["全部", "新概念", "小学教材", "初中教材", "高中教材", "大学四六级", "雅思托福", "英文名著", "外刊新闻", "冒险悬疑", "科幻奇幻", "浪漫爱情", "历史传记", "童话寓言", "短篇小说", "商业科技", "喜剧戏剧", "影视原著", "课外阅读", "其他"]
                v_title = st.text_input("词库书名 (例如: 中考必背词汇 1-50)")
                v_level = st.selectbox("适用级别", base_categories[1:])
                
                upload_method_v = st.radio("单词录入方式", ["手动粘贴词表", "📂 上传本地单词文档 (txt/docx)"], horizontal=True)
                v_raw = ""
                if upload_method_v == "手动粘贴词表":
                    v_raw = st.text_area("粘贴你要上架的纯英文单词 (用逗号或换行隔开)", height=100)
                else:
                    uploaded_file_v = st.file_uploader("选择仅包含单词的文档", type=["txt", "docx"])
                    if uploaded_file_v:
                        v_raw = extract_text_from_file(uploaded_file_v)
                        st.success(f"✅ 成功提取了 {len(v_raw.split())} 个字符段！")
                
                if st.button("🤖 AI 一键解析并发布", type="primary"):
                    if v_title and v_raw.strip():
                        with st.spinner("AI 正在疯狂撰写解析库，请稍候..."):
                            prompt = f"""作为英语教学专家，批量解析以下单词，并严格返回JSON格式：{{"core_vocabulary": [{{"word": "单词", "phonetic": "音标", "translation": "精简释义", "memory_tip": "一句精简的词根或联想记忆法", "usage_examples": "一个简短的英文例句及中文", "tags": "{v_level}"}}]}}。单词列表：\n{v_raw[:1000]}"""
                            try:
                                res = llm_client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], response_format={"type":"json_object"})
                                parsed_json = res.choices[0].message.content
                                supabase.table('public_library').insert({"title": v_title, "category": "公共词库", "content": parsed_json}).execute()
                                st.success("✅ 词库发布成功！全员可见。"); st.rerun()
                            except Exception as e: st.error("解析失败，请减少单词数量重试。")
                    else: st.warning("请填写名称和录入单词。")

        try:
            pub_vocab_raw = supabase.table('public_library').select('*').eq('category', '公共词库').execute().data
            if pub_vocab_raw:
                vocab_options = [f"📚 {v.get('title')}" for v in pub_vocab_raw]
                selected_v_title = st.selectbox("选择大纲词库", vocab_options, label_visibility="collapsed")
                selected_vocab = pub_vocab_raw[vocab_options.index(selected_v_title)]
                
                try:
                    vocab_json = json.loads(selected_vocab.get('content', '{}')).get('core_vocabulary', [])
                    
                    if st.button("⭐ 将这本词书全部加入我的私人生词本", type="primary", use_container_width=True):
                        with st.spinner("正在导入..."):
                            for v in vocab_json:
                                v['user_id'] = CURRENT_USER_ID
                                supabase.table('vocabulary').insert(v).execute()
                            st.success("✅ 导入成功！快去【我的私人生词本】复习吧！")
                    
                    # 🌟 核心：通过 pandas 转为 dataframe 后调用沙盒渲染，发音完美恢复
                    df_pub = pd.DataFrame(vocab_json)
                    render_vocabulary_table(df_pub)
                except: st.error("词库格式异常。")
            else:
                st.info("🌍 馆长还没上传过大纲词汇，敬请期待！")
        except: pass